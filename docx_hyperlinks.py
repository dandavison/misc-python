import os  # Import the os module for path manipulation
import sys

import docx


def extract_and_create_new_doc_with_links(input_docx_path, output_docx_path):
    """
    Extracts hyperlinks from an input .docx file and writes them as plain text
    into a new .docx file.

    Args:
        input_docx_path (str): The path to the source Word document.
        output_docx_path (str): The path where the new Word document
                                will be saved.
    """
    try:
        # 1. Open the source Word document
        source_document = docx.Document(input_docx_path)
    except FileNotFoundError:
        print(f"Error: Input document not found at '{input_docx_path}'")
        return
    except Exception as e:
        print(f"Error opening input document '{input_docx_path}': {e}")
        return

    extracted_urls = []

    # Iterate through all paragraphs and runs to find hyperlinks
    for paragraph in source_document.paragraphs:
        for run in paragraph.runs:
            # Check if the run object is part of a hyperlink field
            # The _element is the underlying lxml element
            if (
                run._element.tag
                == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
            ):
                # Get the relationship ID (r:id) which links to the actual URL
                r_id = run._element.get(
                    "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id"
                )
                if r_id:
                    try:
                        # Access the relationship object to get the target URL
                        # document.part.rels is a dictionary mapping r_id to Relationship objects
                        relationship = source_document.part.rels[r_id]
                        extracted_urls.append(relationship.target_ref)
                    except KeyError:
                        # This shouldn't happen often, but good to handle if a relationship ID is broken
                        print(
                            f"Warning: Broken hyperlink relationship found for r:id '{r_id}'. Skipping."
                        )

    # 3. Create a new Word document
    new_document = docx.Document()

    if extracted_urls:
        new_document.add_heading("Extracted Hyperlinks", level=1)
        new_document.add_paragraph(
            f"Links extracted from: {os.path.basename(input_docx_path)}"
        )
        new_document.add_paragraph("")  # Add a blank line for spacing

        # 4. Write each extracted URL as a plain text paragraph
        for url in extracted_urls:
            new_document.add_paragraph(url)
    else:
        new_document.add_heading("No Hyperlinks Found", level=1)
        new_document.add_paragraph(
            f"No explicit hyperlinks were found in '{os.path.basename(input_docx_path)}'."
        )

    # 5. Save the new document
    try:
        new_document.save(output_docx_path)
        print(f"Successfully created '{output_docx_path}' with extracted hyperlinks.")
    except Exception as e:
        print(f"Error saving new document to '{output_docx_path}': {e}")


if __name__ == "__main__":
    # Ensure proper number of arguments are provided
    if len(sys.argv) < 3:
        print(
            "Usage: python create_links_doc.py <path_to_input_document.docx> <path_to_output_document.docx>"
        )
        sys.exit(1)

    input_doc_path = sys.argv[1]
    output_doc_path = sys.argv[2]

    extract_and_create_new_doc_with_links(input_doc_path, output_doc_path)
